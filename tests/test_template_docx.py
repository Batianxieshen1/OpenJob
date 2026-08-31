"""模板渲染 v9「只换正文」测试：固定区零接触，可变区只换文字，结构不增删。"""

import base64
import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from docx import Document

NL = chr(10)

# 1x1 PNG（照片占位，验证渲染后图片保留）
_PNG_1PX = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR42mNkYPhfDwAChwGA60e6kgAAAABJRU5ErkJggg=="
)


def _add_photo(cell):
    run = cell.paragraphs[0].add_run()
    run.add_picture(BytesIO(_PNG_1PX))


def _build_template(path: Path) -> None:
    """构造 030 蓝标题底稿的骨架：信息表(含照片) + 栏目标题 + bullet + 项目行。"""
    doc = Document()
    info = doc.add_table(rows=2, cols=3)
    info.cell(0, 0).text = "王小明"
    _add_photo(info.cell(0, 1))
    info.cell(1, 0).text = "电话："
    info.cell(1, 1).text = "(+86) 136-7609-7998"
    info.cell(1, 2).text = "candidate@example.com"

    title = doc.add_table(rows=1, cols=1)
    title.cell(0, 0).text = "教育经历"
    intent = doc.add_table(rows=1, cols=1)
    intent.cell(0, 0).text = "求职意向：数据分析（实习）"

    bullet = doc.add_table(rows=1, cols=2)
    bullet.cell(0, 0).text = "•"
    bullet.cell(0, 1).text = "学业成绩：专业前 30%。"

    proj = doc.add_table(rows=1, cols=3)
    proj.cell(0, 0).text = "电商用户增长分析框架（开源）"
    proj.cell(0, 1).text = "独立完成 · GitHub 开源"
    proj.cell(0, 2).text = "2026.05"
    doc.save(str(path))


class TemplateDocxTests(unittest.TestCase):
    def test_replaces_editable_keeps_fixed(self):
        from openjob.ai.resume_engine.template_docx import render_tailored_docx
        from openjob.web.resume_upload import docx_to_markdown

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            _build_template(template)

            tailored = NL.join([
                "王小明",
                "电话： | (+86) 136-7609-7998",
                "candidate@example.com",
                "教育经历",
                "求职意向：AI产品（实习）",
                "• | 学业成绩：专业前 30%，主修数据分析。",
                "电商用户增长分析框架（开源） | 独立完成 | 2026.05",
            ])
            out = Path(tmp) / "tailored.docx"
            render_tailored_docx(template, tailored, out)

            text = docx_to_markdown(out.read_bytes())
            # 可变区已换
            self.assertIn("AI产品（实习）", text)
            self.assertIn("主修数据分析", text)
            # 固定区零接触：姓名/电话/邮箱/栏目标题/'•' 圆点
            self.assertIn("王小明", text)
            self.assertIn("(+86) 136-7609-7998", text)
            self.assertIn("candidate@example.com", text)
            self.assertIn("教育经历", text)
            # 图片保留（照片不动）
            doc = Document(str(out))
            self.assertEqual(len(doc.inline_shapes), 1)
            # 结构不增删：5 张表、表格行数一致
            self.assertEqual(len(doc.tables), 5)
            self.assertEqual([len(t.rows) for t in doc.tables], [2, 1, 1, 1, 1])

    def test_fixed_lines_resist_tampering(self):
        """即使定制稿恶意改写固定区，渲染结果仍保持模板原样。"""
        from openjob.ai.resume_engine.template_docx import render_tailored_docx
        from openjob.web.resume_upload import docx_to_markdown

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            _build_template(template)

            tailored = NL.join([
                "张三",
                "电话： | 13900000000",
                "hacker@example.com",
                "工作经历",
                "求职意向：数据分析（实习）",
                "• | 学业成绩：专业前 5%。",
                "假项目 | 独立完成 | 2027.01",
            ])
            out = Path(tmp) / "tailored.docx"
            render_tailored_docx(template, tailored, out)

            text = docx_to_markdown(out.read_bytes())
            self.assertIn("王小明", text)
            self.assertNotIn("张三", text)
            self.assertIn("(+86) 136-7609-7998", text)
            self.assertNotIn("13900000000", text)
            self.assertNotIn("hacker@example.com", text)
            self.assertIn("教育经历", text)
            self.assertNotIn("工作经历", text)
            # 圆点格不被正文污染
            doc = Document(str(out))
            bullet_cell = doc.tables[3].cell(0, 0).text.strip()
            self.assertEqual(bullet_cell, "•")

    def test_line_count_mismatch_falls_back(self):
        """行数不一致时 difflib 兜底：可对上的行照写，固定区仍不破坏。"""
        from openjob.ai.resume_engine.template_docx import render_tailored_docx
        from openjob.web.resume_upload import docx_to_markdown

        with tempfile.TemporaryDirectory() as tmp:
            template = Path(tmp) / "template.docx"
            _build_template(template)

            tailored = NL.join([
                "王小明",
                "电话： | (+86) 136-7609-7998",
                "candidate@example.com",
                "教育经历",
                "求职意向：数据运营（实习）",
                "• | 学业成绩：专业前 30%。",
                "电商用户增长分析框架（开源） | 独立完成 | 2026.05",
                "多余的一行没有落点",
            ])
            out = Path(tmp) / "tailored.docx"
            render_tailored_docx(template, tailored, out)

            text = docx_to_markdown(out.read_bytes())
            self.assertIn("数据运营（实习）", text)
            self.assertIn("(+86) 136-7609-7998", text)
            doc = Document(str(out))
            self.assertEqual(len(doc.tables), 5)  # 结构不增删


if __name__ == "__main__":
    unittest.main()
