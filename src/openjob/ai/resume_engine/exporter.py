"""导出工具：Markdown 落盘 + PDF 渲染委托。

PDF 渲染复用内置 _render_pdf（Chrome CDP Page.printToPDF 优先，
xhtml2pdf 降级），避免双轨实现；Playwright 打印方案不再保留。
"""

from pathlib import Path

from openjob.ai.resume import _render_pdf as render_pdf
from openjob.ai.resume import _pdf_page_count

__all__ = ["save_markdown", "md_to_pdf", "render_pdf", "_pdf_page_count"]


def save_markdown(text: str, out_path: Path) -> Path:
    out_path.parent.mkdir(parents=True, exist_ok=True)
    out_path.write_text(text, encoding="utf-8")
    return out_path


def md_to_pdf(md_text: str, out_pdf: Path) -> Path | None:
    """Markdown → PDF（CDP 优先，xhtml2pdf 降级）。失败返回 None。"""
    if render_pdf(md_text, out_pdf):
        return out_pdf
    return None


def md_to_docx(md_text: str, out_path: Path) -> Path:
    """把定制简历 Markdown 渲染成 Word 文档（固定专业排版）。

    结构映射：# 姓名/大标题 → Heading1，## 栏目 → Heading2，### 子项 → Heading3，
    "- " → 项目符号段落，**加粗** → 加粗 run。中文字号 10.5pt，A4 页边距 2cm。
    """
    import re

    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(1.8)
        section.left_margin = section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "微软雅黑"
    normal.font.size = Pt(10.5)
    for style_name, size in (("Heading 1", 16), ("Heading 2", 12.5), ("Heading 3", 11)):
        style = doc.styles[style_name]
        style.font.name = "微软雅黑"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = None if hasattr(style.font.color, "rgb") else None

    def add_runs(paragraph, text: str):
        for segment in re.split(r"(\*\*[^*]+\*\*)", text):
            if not segment:
                continue
            if segment.startswith("**") and segment.endswith("**"):
                paragraph.add_run(segment[2:-2]).bold = True
            else:
                paragraph.add_run(segment)

    for raw_line in md_text.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            heading = doc.add_heading(stripped[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            add_runs(paragraph, stripped[2:])
        elif stripped == "---":
            continue
        else:
            paragraph = doc.add_paragraph()
            add_runs(paragraph, stripped)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path
